import requests #imported requests in order to use json/api data
import docx #imported docx to create a word document
from PIL import Image, ImageDraw, ImageFont #imported image, draw, and font from pillow

"""In this block of code I used the same url to create 3 variables. Each variable will generate 
a different random taco recipe."""
taco_recipe_one = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()
taco_recipe_two = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()
taco_recipe_three = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()


image = Image.open('christine-siracusa-vzX2rgUbQXM-unsplash.jpg') #creates object/ opens image downloaded from unsplash
img_draw = ImageDraw.Draw(image) #creates another object, in order to draw/put text on the image
font = ImageFont.truetype('DejaVuSans.ttf', 300) #font variable containing DejavuSans truetype font
img_draw.text([250, 1500], 'Random Taco Cookbook', fill='darkslateblue', font=font)
#specifies text location on image, color, font. Text on image: Random Taco Cookbook
image.thumbnail((800, 800)) #resized image while keeping aspect ratio to approximately 800/800 pixels
image.show() #shows image without having to save
image.save('random_taco_cookbook.jpg') #saves image

document = docx.Document() #created document object
document.add_paragraph('Random Taco Cookbook', 'Title') #adds paragraph with the title style
document.add_picture('random_taco_cookbook.jpg', width=docx.shared.Inches(4), height=docx.shared.Inches(6))
#adds picture saved above to document, width of 4 in, height of 6 in
document.add_paragraph('Credits', 'Heading 1') #new paragraph with heading style
document.add_paragraph('Image: Christine Siracusa on Unsplash') #paragraph containing name of photographer of pic used
document.add_paragraph('Recipes from: https://taco-1150.herokuapp.com/random/?full_taco=true') #random taco recipe API
document.add_paragraph('Code by: Quinn Daniel') #new paragraph, code written by me
document.add_page_break() #adds page break

"""This next block of code defines a function, taco, that I used to make it easier to add 3 random taco recipes to the
document. The function only takes one parameter-recipe. the variables within the function are the names of each key in
the random taco json. I only needed the value of the name key and the recipe key from each of the key's nested 
dictrionaries. This function saves the value for the name key, and the recipe key in each nested dictionary.
This function also adds each of these newly made variables into the document using data from the json. The data will
only be added to the document if the function is called with any of the taco_recipe variables as the parameter."""

def taco(recipe):
    base_layer = recipe['base_layer']['name']
    seasoning = recipe['seasoning']['name']
    mixin = recipe['mixin']['name']
    condiment = recipe['condiment']['name']
    shell = recipe['shell']['name']
    base_layer_recipe = recipe['base_layer']['recipe']
    seasoning_recipe = recipe['seasoning']['recipe']
    mixin_recipe = recipe['mixin']['recipe']
    condiment_recip = recipe['condiment']['recipe']
    shell_recipe = recipe['shell']['recipe']
    document.add_paragraph(f'{base_layer} with {mixin}, {seasoning} and {condiment} in {shell}', 'Title')
    document.add_paragraph(f'{base_layer}', 'Heading 1')
    document.add_paragraph(f'{base_layer_recipe}')
    document.add_paragraph(f'{mixin}', 'Heading 1')
    document.add_paragraph(f'{mixin_recipe}')
    document.add_paragraph(f'{seasoning}', 'Heading 1')
    document.add_paragraph(f'{seasoning_recipe}')
    document.add_paragraph(f'{condiment}', 'Heading 1')
    document.add_paragraph(f'{condiment_recip}')
    document.add_paragraph(f'{shell}', 'Heading 1')
    document.add_paragraph(f'{shell_recipe}')
    document.add_page_break()
    return
taco(taco_recipe_one) #taco function called, generates 1st random recipe
taco(taco_recipe_two) #generates 2nd random recipe
taco(taco_recipe_three) #generates 3rd random recipe

document.save('random_taco_test.docx') #saves document with title: random_taco_test.docx











